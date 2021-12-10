from bs4 import BeautifulSoup
from urllib import parse
from win32com import client as wc
import urllib
import pdfplumber
import docx
import requests.packages.urllib3
import re
import pandas as pd
import numpy
import os
import sys
import time
import traceback
import logging
import xlrd
from tqdm import tqdm

requests.packages.urllib3.disable_warnings()
FORMAT = '%(asctime)s %(levelname)s: %(message)s'
logging.basicConfig(level=logging.ERROR, filename='Databreach.log', filemode='a', format=FORMAT)


class Error_Message:
    def __init__(self, detail):
        self.detail = detail

    """
    def errmsg(self):
        error_class = self.detail.__class__.__name__
        detail_msg = self.detail.args[0]
        log_msg = error_class + ": " + detail_msg
        print(log_msg)
        logging.error(log_msg)
        sys.exit()
    """

    # 錯誤訊息輸出格式
    def errmsg_a(self):
        error_class = self.detail.__class__.__name__  # 取得錯誤類型
        detail = self.detail.args[0]  # 取得詳細內容
        cl, exc, tb = sys.exc_info()  # 取得Call Stack
        lastCallStack = traceback.extract_tb(tb)[-1]  # 取得Call Stack的最後一筆資料
        fileName = lastCallStack[0]  # 取得發生的檔案名稱
        lineNum = lastCallStack[1]  # 取得發生的行號
        funcName = lastCallStack[2]  # 取得發生的函數名稱
        errMsg = "File \"{}\", line {}, in {}: [{}] {}".format(fileName, lineNum, funcName, error_class, detail)
        print(errMsg)
        logging.error(errMsg)
        sys.exit()


class Input_Check:
    def __init__(self, input_addr, condition_1, condition_2, condition_3):
        self.input_addr = input_addr
        self.condition_1 = condition_1
        self.condition_2 = condition_2
        self.condition_3 = condition_3
        self.condition_arr = []

    '''
        檢查輸入的網址是否合法
    '''
    def addr_check(self):
        try:
            r = requests.get(self.input_addr, timeout=3)
            r.raise_for_status()
        except requests.exceptions.HTTPError as errh:
            err = Error_Message(errh)
            err.errmsg_a()
        except requests.exceptions.ConnectionError as errc:
            err = Error_Message(errc)
            err.errmsg_a()
        except requests.exceptions.Timeout as errt:
            err = Error_Message(errt)
            err.errmsg_a()
        except requests.exceptions.RequestException as erro:
            err = Error_Message(erro)
            err.errmsg_a()
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()

    '''
         檢查 外洩條件 是否合法
    '''
    def condition_1_check(self):
        try:
            self.condition_1 = int(self.condition_1)
            self.condition_1 = self.condition_1 + 1
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
        if self.condition_1 > 5:
            print("輸入條件_1超出範圍")
            logging.error('condition_1 out of range')
            sys.exit(-1)
        elif self.condition_1 < 2:
            print("輸入條件_1超出範圍")
            logging.error('condition_1 out of range')
            sys.exit(-1)
        else:
            pass
        return self.condition_1

    '''
         檢查 外洩條件個數 是否合法
    '''
    def condition_2_check(self):
        try:
            self.condition_2 = int(self.condition_2)
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
        if self.condition_2 < 1:
            print("輸入條件_2超出範圍")
            logging.error('condition_2 out of range')
            sys.exit(-1)
        else:
            pass
        return self.condition_2

    '''
         檢查 掃瞄範圍條件 是否合法
    '''
    def condition_3_check(self):
        try:
            self.condition_3 = int(self.condition_3)
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
        if self.condition_3 == 2 or self.condition_3 == 1:
            pass
        else:
            print("輸入條件_3超出範圍")
            logging.error('condition_3 out of range')
            sys.exit(-1)
        return self.condition_3

    '''
         檢查 自訂搜尋範圍 是否合法
    '''
    def condition_4_check(self):
        try:
            if self.condition_3 == 1:
                var = 1
                print("請輸入自定義搜尋範圍之URL關鍵字：(例如:docs.google)\n"
                      "再按一次Enter結束輸入，無輸入時則套用預設全部搜尋")
                while var == 1:
                    tmp = input()
                    if tmp == "":
                        break
                    else:
                        self.condition_arr.append(tmp)
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
        return self.condition_arr


class Obtain_Page:
    def __init__(self, input_addr, condition_3, condition_arr):
        self.input_addr = input_addr
        self.all_url = []
        self.file_name = ""
        self.reorganize_front_page_url = ""
        self.condition_3 = condition_3
        self.condition_arr = condition_arr

    '''
         輸入網址前處理
    '''
    def front_page_url(self):
        fp_url_split = self.input_addr.split('/')
        self.reorganize_front_page_url = fp_url_split[0] + '//' + fp_url_split[2] + '/'
        self.file_name = fp_url_split[2]
        return self.reorganize_front_page_url

    '''
         取回所有首頁中的子網頁超鏈結
         將超鏈結中含有 http 的字串取出
    '''
    def get_url(self):
        self.all_url.append(self.input_addr)
        r = requests.get(self.input_addr, timeout=3)
        r.encoding = 'UTF8'
        soup = BeautifulSoup(r.text, 'lxml')
        a_tags = soup.find_all('a', href=True)
        tmp_arr = []
        for tag in a_tags:
            tag = tag.get('href')
            # print(tag)
            if tag =="":
                continue
            elif '#' in tag:
                continue
            elif 'javascript' in tag:
                continue
            elif '@' in tag:
                continue
            elif tag[0] == "/":
                try:
                    if tag[1] == "/":
                        tag_s = "https:" + tag
                        self.all_url.append(tag_s)
                    else:
                        tag = tag[1:]
                except:
                    pass
            elif 'http' in tag and self.condition_3 == 2:
                continue
            elif 'http' in tag and self.condition_3 == 1:
                tmp_arr.append(tag)
            elif '?' in tag:
                tag = self.reorganize_front_page_url + tag
                self.all_url.append(tag)
            else:
                tag = self.reorganize_front_page_url + urllib.parse.quote(tag)
                self.all_url.append(tag)
        if len(self.condition_arr) > 0:
            for i in range(len(self.condition_arr)):
                for j in range(len(tmp_arr)):
                    if self.condition_arr[i] in tmp_arr[j]:
                        self.all_url.append(tmp_arr[j])
        else:
            self.all_url.extend(tmp_arr)
        all_url = numpy.unique(self.all_url)
        print("掃描進行中...")
        time.sleep(3)
        return all_url

    def mkdir_file(self):
        try:
            path = './' + self.file_name + '/'
            if not os.path.isdir(path):
                os.mkdir(path)
            return path
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
            print("建立資料夾失敗")
            sys.exit(-1)


class Check_Formula:
    def __init__(self):
        pass

    '''
         身分證字號格式檢查(判別式)
    '''
    def id_num_check(self, id_num):
        try:
            id2 = []
            for j in range(len(id_num)):
                sp0 = id_num[j]
                sp1 = repr(sp0)
                sp2 = (sp1[1:len(sp1) - 1])
                s = list(sp2)
                sum_id = 0
                if s[0] == 'A':
                    sum_id = sum_id + 1
                elif s[0] == 'B':
                    sum_id = sum_id + 10
                elif s[0] == 'C':
                    sum_id = sum_id + 19
                elif s[0] == 'D':
                    sum_id = sum_id + 28
                elif s[0] == 'E':
                    sum_id = sum_id + 37
                elif s[0] == 'F':
                    sum_id = sum_id + 46
                elif s[0] == 'G':
                    sum_id = sum_id + 55
                elif s[0] == 'H':
                    sum_id = sum_id + 64
                elif s[0] == 'I':
                    sum_id = sum_id + 39
                elif s[0] == 'J':
                    sum_id = sum_id + 73
                elif s[0] == 'K':
                    sum_id = sum_id + 82
                elif s[0] == 'L':
                    sum_id = sum_id + 2
                elif s[0] == 'M':
                    sum_id = sum_id + 11
                elif s[0] == 'N':
                    sum_id = sum_id + 20
                elif s[0] == 'O':
                    sum_id = sum_id + 48
                elif s[0] == 'P':
                    sum_id = sum_id + 29
                elif s[0] == 'Q':
                    sum_id = sum_id + 38
                elif s[0] == 'R':
                    sum_id = sum_id + 47
                elif s[0] == 'S':
                    sum_id = sum_id + 56
                elif s[0] == 'T':
                    sum_id = sum_id + 65
                elif s[0] == 'U':
                    sum_id = sum_id + 74
                elif s[0] == 'V':
                    sum_id = sum_id + 83
                elif s[0] == 'W':
                    sum_id = sum_id + 21
                elif s[0] == 'X':
                    sum_id = sum_id + 3
                elif s[0] == 'Y':
                    sum_id = sum_id + 12
                elif s[0] == 'Z':
                    sum_id = sum_id + 30
                else:
                    sum_id = sum_id + 0
                # 身分證字號判別式
                sum1 = 8 * int(s[1]) + 7 * int(s[2]) + 6 * int(s[3]) + 5 * int(s[4]) + 4 * int(s[5]) + 3 * int(
                    s[6]) + 2 * int(
                    s[7]) + 1 * int(s[8]) + int(s[9])
                sum_id = sum_id + sum1
                sum_id = sum_id % 10
                # 性別欄位判別
                gender = 0
                if int(s[1]) == 1:
                    gender = 1
                elif int(s[1]) == 2:
                    gender = 1
                else:
                    pass
                # 必須符合兩個條件
                if sum_id == 0 and gender == 1:
                    id2.append(sp0)
                else:
                    pass
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
            print("身分證字號檢查時發生錯誤")
            sys.exit(-1)
        return id2

    '''
         電話號碼格式檢查
    '''
    def phone_check(self, phone):
        try:
            ph_2 = []
            for k in range(len(phone)):
                rp = repr(phone[k])
                rp = rp[1:-1]

                if rp[0].isdigit():
                    pass
                else:
                    rp = rp[1:]

                if rp[-1].isdigit():
                    pass
                else:
                    rp = rp[:-2]

                if len(rp) < 8:
                    pass
                else:
                    ph_2.append(rp)

        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
            print("電話號碼檢查時發生錯誤")
            sys.exit(-1)
        return ph_2

    '''
         檢查 黑名單 (根據 filter.txt)
    '''
    def name_check_b(self, name):
        try:
            name_filter = []
            f = open('filter.txt', 'r', encoding="utf-8")
            for line in f.readlines():
                name_filter.append(line[:-1])
            f.close
            name_2 = [x for x in name if x not in name_filter]
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
            print("黑名單檢查時發生錯誤")
            sys.exit(-1)
        return name_2, len(name), len(name_2)

    '''
         檢查 白名單 (自訂義)
    '''
    def name_check_w(self, name):
        try:
            name_allow = []
            f = open('allow.txt', 'r', encoding="utf-8")
            for line in f.readlines():
                name_allow.append(line[:-1])
            f.close
            name_1 = [x for x in name if x not in name_allow]
        except Exception as e:
            err = Error_Message(e)
            err.errmsg_a()
            print("白名單檢查時發生錯誤")
            sys.exit(-1)
        return name_1


class Analysis:
    def __init__(self):
        self.a = ""
        self.all_url = None
        self.test = None
        self.path = ""
        self.err = ""
        self.fullText = ""
        self.tmp_array = []
        self.output_array = []
        self.suc = 0
        self.flag = 0
        # self.r_name_check = re.compile("姓名")
        self.r_name = re.compile(
            r"([陳林黃張李王吳劉蔡楊許鄭謝洪郭邱曾廖賴徐周葉蘇莊呂江何蕭羅高潘簡朱鍾游彭詹胡施沈余盧梁趙顏柯翁魏孫戴范方宋鄧杜傅侯曹薛丁卓阮馬董温唐藍石蔣古紀姚連馮歐程湯黄田康姜白汪鄒尤巫鐘黎涂龔嚴韓袁金童陸夏柳凃邵錢伍倪溫于譚駱熊任甘秦顧毛章史官萬俞雷粘]{1})"
            r"([\u4E00-\u9FFF]{2})")
        '''
        self.r_name = re.compile(
            r"(\s|:|：?)([陳林黃張李王吳劉蔡楊許鄭謝洪郭邱曾廖賴徐周葉蘇莊呂江何蕭羅高潘簡朱鍾游彭詹胡施沈余盧梁趙顏柯翁魏孫戴范方宋鄧杜傅侯曹薛丁卓阮馬董温唐藍石蔣古紀姚連馮歐程湯黄田康姜白汪鄒尤巫鐘黎涂龔嚴韓袁金童陸夏柳凃邵錢伍倪溫于譚駱熊任甘秦顧毛章史官萬俞雷粘]{1})"
            r"([\u4E00-\u9FFF]{2})(\s{1}?)")
        '''
        # \u2E80-\u9FFF
        self.r_id_num = re.compile(r"[a-z]\d{9}|[A-Z]\d{9}")
        # self.r3 = re.compile(r"(0{1})(\d{1,3})(-{1})(\d{5,8})")
        self.r_phone_len_1 = re.compile(r"(\d+)(-)(\d+)(-?)(\d*)")
        self.r_phone_len_2 = re.compile(r"([(])(\d+)([)])(\s?)(\d+)(-?)(\d*)")
        self.r_phone_1_1 = re.compile(r"(886|0)(-?)([3-8])(-)(\d{3})(-?)(\d{4})")
        self.r_phone_1_2 = re.compile(r"(886|0)(-?)(2)(-)(\d{4})(-?)(\d{4})")
        self.r_phone_1_3 = re.compile(r"(886|0)(-?)(37)(-)(\d{2})(-?)(\d{4})")
        self.r_phone_1_4 = re.compile(r"(886|0)(-?)(49)(-)(\d{3})(-?)(\d{4})")
        self.r_phone_1_5 = re.compile(r"(886|0)(-?)(82)(-)(\d{2})(-?)(\d{4})")
        self.r_phone_1_6 = re.compile(r"(886|0)(-?)(89)(-)(\d{2})(-?)(\d{4})")
        self.r_phone_1_7 = re.compile(r"(886|0)(-?)(836)(-)(\d{1})(-?)(\d{4})")
        self.r_phone_1_8 = re.compile(r"(886|0)(-?)(9\d{2})(-)(\d{3})(-?)(\d{3})")
        self.r_phone_1_9 = re.compile(r"(886|0)(-?)(8\d{2})(-)(\d{3})(-?)(\d{3})")
        self.r_phone_2_1 = re.compile(r"([(])(0[3-8])([)])(\s?)(\d{3})(-?)(\d{4})")
        self.r_phone_2_2 = re.compile(r"([(])(02)([)])(\s?)(\d{4})(-?)(\d{4})")
        self.r_phone_2_3 = re.compile(r"([(])(037)([)])(\s?)(\d{2})(-?)(\d{4})")
        self.r_phone_2_4 = re.compile(r"([(])(049)([)])(\s?)(\d{3})(-?)(\d{4})")
        self.r_phone_2_5 = re.compile(r"([(])(082)([)])(\s?)(\d{2})(-?)(\d{4})")
        self.r_phone_2_6 = re.compile(r"([(])(089)([)])(\s?)(\d{2})(-?)(\d{4})")
        self.r_phone_2_7 = re.compile(r"([(])(0836)([)])(\s?)(\d{1})(-?)(\d{4})")
        self.r_addr = re.compile(r"(台北|臺北|新北|桃園|台中|臺中|台南|臺南|高雄|苗栗|彰化|南投|雲林|嘉義|屏東|宜蘭|花蓮|台東|臺東|澎湖|金門|連江|基隆|新竹+)"
                             r"([\u4E00-\u9FFF]+)(鄉|鎮|市|區?)([\u4E00-\u9FFF]+)(街|大道|路+)([一-十]?段?)([一-十]|百|千|\d+)(-?)(\w*)(號+)")

    '''
         將所有子網頁取回，並將其儲存至目標資料夾
         若是該超鏈結沒有回應則在輸出報表中該列打勾
    '''
    def document(self, i, all_url, path):
        self.err = ""
        self.all_url = all_url
        self.path = path
        url = str(self.all_url)
        try:
            self.test = requests.get(self.all_url, verify=False, timeout=5)
        except:
            self.err = "V"
            return url, self.err, self.suc
        try:
            if '.pdf' in url:
                # 下載檔案
                writefile = path + str(i) + ".pdf"
                with open(writefile, 'wb') as f:
                    f.write(self.test.content)
                f.close()
                # 解析
                pdf = pdfplumber.open(writefile)
                for j in range(len(pdf.pages)):
                    p = pdf.pages[j]
                    text = p.extract_text()
                    self.fullText = self.fullText + str(text)
            elif '.docx' in url:
                writefile = path + str(i) + ".docx"
                with open(writefile, 'wb') as f:
                    f.write(self.test.content)
                f.close()
                # 解析
                doc = docx.Document(writefile)  # io.BytesIO(test.content)
                for para in doc.paragraphs:
                    self.fullText = self.fullText + para.text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self.fullText = self.fullText + para.text
                self.flag = 1
            elif '.doc' in url and self.flag == 0:
                # 下載檔案
                writefile = path + str(i) + ".doc"
                with open(writefile, 'wb') as f:
                    f.write(self.test.content)
                f.close()
                # 轉檔
                location = os.getcwd()
                word = wc.Dispatch('Word.Application')
                file_location = str(location) + path + str(i) + ".doc"
                readfile = str(location) + path + str(i) + ".docx"
                doc = word.Documents.Open(file_location)
                doc.SaveAs(readfile, 12, False, "", True, "", False, False, False, False)
                doc.Close()
                word.Quit()
                os.remove(file_location)
                # 解析
                doc = docx.Document(readfile)  # io.BytesIO(test.content)
                for para in doc.paragraphs:
                    self.fullText = self.fullText + para.text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self.fullText = self.fullText + para.text
            elif '.xls' in url:
                # 下載檔案
                writefile = path + str(i) + ".xls"
                with open(writefile, 'wb') as f:
                    f.write(self.test.content)
                f.close()
                data = xlrd.open_workbook(writefile,logfile="Databreach.log")
                for sheet in data.sheets():
                    for i in range(sheet.nrows):
                        for j in range(sheet.ncols):
                            text = sheet.cell_value(i, j)
                            stext = str(text)
                            self.fullText = self.fullText + stext

            else:
                # 下載檔案
                writefile = self.path + str(i) + ".html"
                with open(writefile, 'wb') as f:
                    f.write(self.test.content)
                f.close()
                # 解析
                self.test.encoding = 'UTF8'
                soup_leaf = BeautifulSoup(self.test.text, 'lxml')
                self.fullText = soup_leaf.text
            self.suc = self.suc + 1
            self.flag = 0
        except:
            self.err = "V"
            return url, self.err, self.suc
        return url, self.err, self.suc

    '''
         將每個子網頁過 regular expression
         目前包含四種檢出目標，人名、身份證字號、電話號碼、地址
         並且會將重複的檢出結果去除
    '''
    def reg_find(self):
        name_tmp = ""
        name_out = []
        # name_check = re.findall(self.r_name_check, self.fullText)
        name = re.findall(self.r_name, self.fullText)
        # print(name_check)
        # print(len(name_check))
        # if len(name_check) > 0:
        for k in range(len(name)):
            '''
            name[k] = name[k][1:-1]
            for m in range(len(name[k])):                
                name_tmp = name_tmp + str(name[k][m])
            '''
            if "先生" in name[k][1]:
                continue
            elif "小姐" in name[k][1]:
                continue
            else:
                name_tmp = str(name[k][0]) + str(name[k][1])
                name_out.append(name_tmp)
            name_tmp = ""
        id_num = re.findall(self.r_id_num, self.fullText)
        addr = re.findall(self.r_addr, self.fullText)
        phone_len_tmp = ""
        phone_len_tmp2 = ""
        phone_len_tmp3 = ""
        phone_len_tmp4 = ""
        phone_out = []
        phone_len1 = re.findall(self.r_phone_len_1, self.fullText)
        for i in range(len(phone_len1)):
            for j in range(len(phone_len1[i])):
                phone_len_tmp = phone_len_tmp + str(phone_len1[i][j])
            if len(phone_len_tmp) > 15:
                phone_len_tmp = ""
            phone_f1 = re.findall(self.r_phone_1_1, phone_len_tmp)
            phone_f2 = re.findall(self.r_phone_1_2, phone_len_tmp)
            phone_f3 = re.findall(self.r_phone_1_3, phone_len_tmp)
            phone_f4 = re.findall(self.r_phone_1_4, phone_len_tmp)
            phone_f5 = re.findall(self.r_phone_1_5, phone_len_tmp)
            phone_f6 = re.findall(self.r_phone_1_6, phone_len_tmp)
            phone_f7 = re.findall(self.r_phone_1_7, phone_len_tmp)
            phone_f8 = re.findall(self.r_phone_1_8, phone_len_tmp)
            phone_f9 = re.findall(self.r_phone_1_9, phone_len_tmp)
            phone_len_tmp = ""
            phone_f1.extend(phone_f2)
            phone_f1.extend(phone_f3)
            phone_f1.extend(phone_f4)
            phone_f1.extend(phone_f5)
            phone_f1.extend(phone_f6)
            phone_f1.extend(phone_f7)
            phone_f1.extend(phone_f8)
            phone_f1.extend(phone_f9)
            if phone_f1 == []:
                continue
            for k in range(len(phone_f1)):
                for m in range(len(phone_f1[k])):
                    phone_len_tmp2 = phone_len_tmp2 + str(phone_f1[k][m])
                phone_out.append(phone_len_tmp2)
                phone_len_tmp2 = ""

        phone_len2 = re.findall(self.r_phone_len_2, self.fullText)
        for i in range(len(phone_len2)):
            for j in range(len(phone_len2[i])):
                phone_len_tmp3 = phone_len_tmp3 + str(phone_len2[i][j])
            if len(phone_len_tmp3) > 15:
                phone_len_tmp3 = ""
            phone_f9 = re.findall(self.r_phone_2_1, phone_len_tmp3)
            phone_f10 = re.findall(self.r_phone_2_2, phone_len_tmp3)
            phone_f11 = re.findall(self.r_phone_2_3, phone_len_tmp3)
            phone_f12 = re.findall(self.r_phone_2_4, phone_len_tmp3)
            phone_f13 = re.findall(self.r_phone_2_5, phone_len_tmp3)
            phone_f14 = re.findall(self.r_phone_2_6, phone_len_tmp3)
            phone_f15 = re.findall(self.r_phone_2_7, phone_len_tmp3)
            phone_len_tmp3 = ""
            phone_f9.extend(phone_f10)
            phone_f9.extend(phone_f11)
            phone_f9.extend(phone_f12)
            phone_f9.extend(phone_f13)
            phone_f9.extend(phone_f14)
            phone_f9.extend(phone_f15)
            if phone_f9 == []:
                continue
            for k in range(len(phone_f9)):
                for m in range(len(phone_f9[k])):
                    phone_len_tmp4 = phone_len_tmp4 + str(phone_f9[k][m])
                phone_out.append(phone_len_tmp4)
                phone_len_tmp4 = ""

        addr_2 = []
        addr_tmp = ""
        for i in range(len(addr)):
            for j in range(len(addr[i])):
                addr_tmp = addr_tmp + str(addr[i][j])
            addr_2.append(addr_tmp)
            addr_tmp = ""
        check = Check_Formula()
        name = numpy.unique(name_out)
        # name_1 = check.name_check_w(name)
        name_2, t1, t2 = check.name_check_b(name)
        # name_2.extend(name_1)
        id_2 = check.id_num_check(id_num)
        id_2 = numpy.unique(id_2)
        phone_out = numpy.unique(phone_out)
        addr_2 = numpy.unique(addr_2)
        return name_2, addr_2, id_2, phone_out, t1, t2
    """
    def name_check(self,name):
        total_count_a = 0
        total_count_b = 0
        name_filter = []
        f = open('qwe.txt','r',encoding="utf-8")
        for line in f.readlines():
            name_filter.append(line[:-1])
        f.close
        total_count_a += len(name)
        name_2 = [x for x in name if x not in name_filter]
        total_count_b += len(name_2)
        return name_2, total_count_a, total_count_b
    """



class Generate_Report:
    def __init__(self):
        self.url = ""
        self.err = ""
        self.name = ""
        self.addr = ""
        self.id_num = ""
        self.phone = ""
        self.table_row = []
        self.table = []
        self.risk = 0

    '''
        將每個子網頁掃描結果輸入報表
    '''
    def generate_table(self, url, err, name, addr, id_num, phone):
        if err == "":
            self.url = url
            self.err = err
            self.name = name
            self.addr = addr
            self.id_num = id_num
            self.phone = phone
        else:
            self.url = url
            self.err = err
            self.name = ""
            self.addr = ""
            self.id_num = ""
            self.phone = ""
        self.table_row.append(self.url)
        self.table_row.append(self.err)
        self.table_row.append(len(self.name))
        self.table_row.append(len(self.id_num))
        self.table_row.append(len(self.phone))
        self.table_row.append(len(self.addr))
        self.table_row.append(self.name)
        self.table_row.append(self.id_num)
        self.table_row.append(self.phone)
        self.table_row.append(self.addr)
        row_tuple = tuple(self.table_row)
        self.table.append(row_tuple)

    '''
        檢查是否符合輸入條件
    '''
    def condition_check(self, condition_1, condition_2):
        if self.table_row[condition_1] >= condition_2:
            self.risk = self.risk + 1
        self.table_row.clear()
        return self.risk

    '''
        將掃描結果生成 csv 檔
    '''
    def wirte_file(self, path):
        file_name = path + "掃描結果.csv"
        tmp4 = pd.DataFrame(self.table,
                            columns=["網址", "連線失敗", "人名個數", "身分證字號個數", "電話號碼個數", "地址個數", "詳細人名", "詳細身分證字號", "詳細電話號碼",
                                     "詳細地址"])
        tmp4.to_csv(file_name, encoding="utf_8-sig")


class Interface:
    def __init__(self):
        pass

    '''
        輸入介面
    '''
    def input_box(self):
        print("請輸入學校網址(例如 https://www.ee.ncku.edu.tw/)：")
        input_addr = input()
        # input_addr =  "http://health.tn.edu.tw/result/file/1620028273_110%E5%B9%B4%E8%87%BA%E5%8D%97%E5%B8%82%E5%9C%8B%E5%B0%8F%E7%B1%83%E7%90%83%E5%B0%8D%E6%8A%97%E8%B3%BD%E6%9A%A8%E6%95%99%E8%82%B2%E9%83%A8%E9%AB%94%E8%82%B2%E7%BD%B2109%E5%AD%B8%E5%B9%B4%E5%BA%A6%E5%9C%8B%E5%B0%8F%E7%B1%83%E7%90%83%E8%81%AF%E8%B3%BD%E9%A0%90%E8%B3%BD%E7%AB%B6%E8%B3%BD%E8%A6%8F%E7%A8%8B.pdf"
        # "https://www.ee.ncku.edu.tw/form.php?type=master"
        # "https://www.ee.ncku.edu.tw"
        print("請輸入判斷個資外洩條件之編號:\n"
              "1.人名個數\n"
              "2.身分證字號個數\n"
              "3.電話號碼個數\n"
              "4.地址個數")
        condition_1 = input()
        print("請輸入要依照多少個數作為條件:")
        condition_2 = input()
        print("請輸入是否掃描非本校之連結:\n"
              "1.是\n"
              "2.否")
        condition_3 = input()
        return input_addr, condition_1, condition_2, condition_3,

    '''
        輸出頁面
    '''
    def output(self, suc, risk):
        print("成功連接 " + str(suc) + " 個網頁")
        print("其中有 " + str(risk) + " 個網頁可能有個資外洩疑慮")
        print("Success")
        pass


def main():
    t = 0
    tt = 0
    io = Interface()
    # 輸入條件
    input_addr, condition_1, condition_2, condition_3 = io.input_box()
    # 輸入條件檢查
    step_1 = Input_Check(input_addr, condition_1, condition_2, condition_3)
    step_1.addr_check()
    condition_1 = step_1.condition_1_check()
    condition_2 = step_1.condition_2_check()
    condition_3 = step_1.condition_3_check()
    condition_arr = step_1.condition_4_check()
    # 取的輸入網頁中的所有子網頁超鏈結
    step_2 = Obtain_Page(input_addr, condition_3, condition_arr)
    step_2.front_page_url()
    path = step_2.mkdir_file()
    all_url = step_2.get_url()
    # 掃描所有子網頁
    step_3 = Analysis()
    step_4 = Generate_Report()
    for i in tqdm(range(len(all_url))):
        # 檢查該子網頁是否有回應
        url, err, suc = step_3.document(i, all_url[i], path)
        # 解析超鏈結內容並檢查是否有匹配字串
        name, addr, id_num, phone, t1, t2 = step_3.reg_find()
        # 原數量人名數量
        t += t1
        # 過黑名單後數量
        tt += t2
        step_4.generate_table(url, err, name, addr, id_num, phone)
        risk = step_4.condition_check(condition_1, condition_2)
    # 生成報表
    step_4.wirte_file(path)
    # 輸出結果
    io.output(suc, risk)

# ------------------------------------------------------------------


if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        err = Error_Message(e)
        err.errmsg_a()